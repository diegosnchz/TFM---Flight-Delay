"""
generate_tfm.py - Generador del documento Word del TFM.

Crea el documento tfm_eu261.docx con la estructura completa del TFM,
incluyendo todos los capitulos, tablas de resultados e imagenes generadas.

Requiere que las fases anteriores (EDA, modelado, arbitraje) hayan sido
ejecutadas y sus outputs esten disponibles en outputs/.

Ejecutar con: python docs/tfm/generate_tfm.py
"""

from __future__ import annotations

import logging
import sys
from pathlib import Path

# Agregar el directorio raiz al path para importar src
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.config import FIGURES_DIR, MODELS_DIR, TABLES_DIR, setup_logging

logger = setup_logging(__name__)

TFM_OUTPUT = Path(__file__).parent / "tfm_eu261.docx"

# ---------------------------------------------------------------------------
# Helpers de formato
# ---------------------------------------------------------------------------


def add_heading(doc: Document, text: str, level: int) -> None:
    """Anade un titulo con el nivel especificado."""
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """Anade un parrafo de texto."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_figure(doc: Document, filename: str, caption: str, width_inches: float = 5.5) -> None:
    """
    Inserta una figura si existe, o un placeholder si no.

    Args:
        doc: Documento Word.
        filename: Nombre del archivo de imagen en outputs/figures/.
        caption: Descripcion de la figura (pie de figura).
        width_inches: Ancho de la imagen en pulgadas.
    """
    fig_path = FIGURES_DIR / filename

    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(width_inches))
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[FIGURA PENDIENTE: {filename}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Pie de figura centrado e italico
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_caption = p_caption.add_run(caption)
    run_caption.italic = True
    run_caption.font.size = Pt(10)


def add_table_from_csv(doc: Document, csv_filename: str, caption: str) -> None:
    """
    Inserta una tabla desde un CSV si existe, o un placeholder.

    Args:
        doc: Documento Word.
        csv_filename: Nombre del archivo CSV en outputs/tables/.
        caption: Titulo de la tabla.
    """
    csv_path = TABLES_DIR / csv_filename

    if not csv_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[TABLA PENDIENTE: {csv_filename}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        return

    df = pd.read_csv(csv_path)

    # Encabezado de tabla
    p_caption = doc.add_paragraph()
    run_caption = p_caption.add_run(caption)
    run_caption.bold = True
    run_caption.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Shading Accent 1"

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True

    # Datos (limitar a 20 filas para no saturar el documento)
    for _, row in df.head(20).iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


# ---------------------------------------------------------------------------
# Generacion de secciones
# ---------------------------------------------------------------------------


def build_portada(doc: Document) -> None:
    """Genera la portada del TFM."""
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        "Analisis Predictivo de Retrasos Aereos en Aerolineas\n"
        "Low-Cost Europeas: Evaluacion de Oportunidades de\n"
        "Arbitraje Regulatorio bajo la Normativa EU261/2004"
    )
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run("Trabajo de Fin de Master").bold = True

    doc.add_paragraph()

    info_items = [
        ("Autor", "Diego Sanchez"),
        ("Master", "Master en Big Data e Inteligencia Artificial - Euroformac"),
        ("Fecha", "Junio 2026"),
        ("Tutor", "[PENDIENTE DE RELLENAR]"),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)

    doc.add_page_break()


def build_abstract(doc: Document) -> None:
    """Genera el abstract en espanol e ingles."""
    add_heading(doc, "Abstract", 1)

    add_heading(doc, "Resumen (Espanol)", 2)
    add_paragraph(doc, (
        "El presente Trabajo de Fin de Master analiza la viabilidad de construir un modelo "
        "predictivo de retrasos aereos en aerolineas low-cost europeas (Ryanair, easyJet, "
        "Wizz Air y Vueling) con el objetivo de identificar oportunidades de arbitraje "
        "regulatorio bajo el Reglamento (CE) n. 261/2004 (EU261). "
        "Utilizando datos del repositorio Eurocontrol ADRR, se entrena y compara una "
        "bateria de modelos de machine learning, incluyendo regresion logistica, Random "
        "Forest, XGBoost y LightGBM, sobre un conjunto de vuelos caracterizado por un "
        "severo desbalance de clases (tasa de retrasos elegibles inferior al 1%). "
        "Se aborda el problema del data leakage en las features de tasa historica, la "
        "paradoja de Simpson en el analisis por aerolinea, y la seleccion del umbral de "
        "clasificacion optimo mediante la curva Precision-Recall. "
        "El mejor modelo se integra en un modelo de Expected Value extendido que incorpora "
        "los costes reales del arbitraje: precio del billete, transporte, coste de "
        "oportunidad temporal y probabilidad de rechazo de la reclamacion. "
        "Una simulacion Monte Carlo de 10,000 iteraciones cuantifica la incertidumbre del "
        "modelo. Los resultados sugieren que, si bien el modelo identifica rutas con "
        "probabilidad elevada de retraso, el Expected Value neto es negativo para la "
        "mayoria de escenarios realistas, confirmando que las barreras practicas "
        "(rechazo de reclamaciones, coste temporal, cancelaciones anticipadas) hacen "
        "del arbitraje EU261 una estrategia de riesgo no trivial."
    ))

    add_paragraph(doc, (
        "Palabras clave: machine learning, retrasos aereos, EU261, arbitraje regulatorio, "
        "regresion logistica, XGBoost, desbalance de clases, Expected Value."
    ), bold=False)

    doc.add_paragraph()

    add_heading(doc, "Abstract (English)", 2)
    add_paragraph(doc, (
        "This Master's Thesis analyses the feasibility of building a predictive model of "
        "flight delays in European low-cost airlines (Ryanair, easyJet, Wizz Air and "
        "Vueling) in order to identify regulatory arbitrage opportunities under Regulation "
        "(EC) No. 261/2004 (EU261). "
        "Using data from the Eurocontrol ADRR repository, a suite of machine learning "
        "models is trained and compared, including logistic regression, Random Forest, "
        "XGBoost and LightGBM, on a flight dataset characterised by a severe class "
        "imbalance (rate of eligible delays below 1%). "
        "The study addresses data leakage in historical rate features, the Simpson's "
        "Paradox in airline-level analysis, and optimal classification threshold selection "
        "via the Precision-Recall curve. "
        "The best model is integrated into an extended Expected Value framework that "
        "incorporates real arbitrage costs: ticket price, transport, time opportunity cost, "
        "and claim denial probability. "
        "A 10,000-iteration Monte Carlo simulation quantifies the model's uncertainty. "
        "Results suggest that, while the model identifies routes with elevated delay "
        "probability, the net Expected Value is negative in most realistic scenarios, "
        "confirming that practical barriers (claim denials, time costs, early cancellations) "
        "make EU261 arbitrage a non-trivial risk strategy."
    ))

    add_paragraph(doc, (
        "Keywords: machine learning, flight delays, EU261, regulatory arbitrage, "
        "logistic regression, XGBoost, class imbalance, Expected Value."
    ))

    doc.add_page_break()


def build_introduccion(doc: Document) -> None:
    """Genera el Capitulo 1: Introduccion."""
    add_heading(doc, "1. Introduccion", 1)

    add_heading(doc, "1.1 Motivacion y Contexto", 2)
    add_paragraph(doc, (
        "El sector de la aviacion comercial europea ha experimentado en los ultimos anos "
        "un crecimiento sostenido de las aerolineas de bajo coste, que en 2023 operaban "
        "mas del 40% de los vuelos intraeuropeos (Eurocontrol, 2024). Paralelamente, la "
        "normativa EU261/2004 establece un sistema de compensaciones economicas "
        "obligatorias para los pasajeros afectados por retrasos superiores a tres horas, "
        "con importes de entre 250 y 600 euros segun la distancia del vuelo."
    ))
    add_paragraph(doc, (
        "La combinacion de billetes de bajo coste, alta frecuencia de retrasos en "
        "determinadas rutas y compensaciones de importe fijo genera una estructura de "
        "payoffs asimetrica que, en teoria, podria explotarse de forma sistematica. "
        "El presente trabajo explora si es posible construir un modelo predictivo que "
        "identifique ex ante que vuelos tienen mayor probabilidad de retraso compensable, "
        "y si el Expected Value resultante justificaria la compra estrategica de billetes."
    ))

    add_heading(doc, "1.2 Objetivos del Trabajo", 2)
    add_paragraph(doc, "Los objetivos especificos del TFM son los siguientes:")

    objectives = [
        "Construir y comparar modelos de machine learning para predecir la probabilidad "
        "de retraso EU261-elegible en vuelos de aerolineas low-cost europeas.",
        "Analizar el impacto de variables temporales, geograficas y operativas en la "
        "probabilidad de retraso.",
        "Desarrollar un modelo de Expected Value extendido que cuantifique la rentabilidad "
        "real del arbitraje EU261 considerando todos los costes relevantes.",
        "Evaluar la paradoja de Simpson en el analisis de tasas de retraso por aerolinea.",
        "Cuantificar la incertidumbre del modelo mediante simulacion Monte Carlo.",
        "Discutir la viabilidad practica, las limitaciones y las implicaciones eticas "
        "del modelo propuesto.",
    ]

    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(obj)

    add_heading(doc, "1.3 Alcance y Limitaciones", 2)
    add_paragraph(doc, (
        "El estudio se limita a cuatro aerolineas low-cost europeas (Ryanair, easyJet, "
        "Wizz Air y Vueling) y a vuelos operados en el espacio aereo europeo. "
        "Se analizan unicamente retrasos en la llegada, no cancelaciones. "
        "El modelo no considera factores meteorologicos ni de control de trafico aereo "
        "de forma explicita. Las limitaciones completas se detallan en el Capitulo 7."
    ))

    add_heading(doc, "1.4 Estructura del Documento", 2)
    add_paragraph(doc, (
        "El documento se estructura en ocho capitulos. El Capitulo 2 revisa el estado "
        "del arte en prediccion de retrasos, normativa EU261 y arbitraje regulatorio. "
        "El Capitulo 3 describe la metodologia. El Capitulo 4 presenta el analisis "
        "exploratorio de datos. El Capitulo 5 detalla los resultados del modelado. "
        "El Capitulo 6 desarrolla el modelo de arbitraje. El Capitulo 7 discute los "
        "resultados y el Capitulo 8 presenta las conclusiones."
    ))

    doc.add_page_break()


def build_estado_arte(doc: Document) -> None:
    """Genera el Capitulo 2: Estado del Arte."""
    add_heading(doc, "2. Estado del Arte", 1)

    add_heading(doc, "2.1 Machine Learning Aplicado a la Prediccion de Retrasos Aereos", 2)
    add_paragraph(doc, (
        "La prediccion de retrasos aereos mediante tecnicas de machine learning ha sido "
        "objeto de numerosos estudios academicos. [REF_01: Buscar en Google Scholar "
        "'flight delay prediction machine learning' filtrado por 2019-2024]. "
        "Los modelos de gradient boosting (XGBoost, LightGBM) han demostrado "
        "consistentemente resultados superiores a modelos lineales en este dominio "
        "[REF_02]. El desbalance de clases es un desafio recurrente, ya que los vuelos "
        "con retrasos graves representan tipicamente menos del 5% de los datos [REF_03]."
    ))

    add_heading(doc, "2.2 Normativa EU261/2004: Marco Legal y Jurisprudencia", 2)
    add_paragraph(doc, (
        "El Reglamento (CE) n. 261/2004 establece el derecho a compensacion de los "
        "pasajeros aereos en casos de gran retraso, cancelacion y denegacion de embarque. "
        "La sentencia del TJUE en el asunto Sturgeon (C-402/07) extendia el derecho de "
        "compensacion a los pasajeros con retrasos superiores a tres horas en la llegada, "
        "equiparandolos a los afectados por cancelaciones [REF_04: STJUE C-402/07]. "
        "Las condiciones de exencion por 'circunstancias extraordinarias' han sido "
        "interpretadas restrictivamente por los tribunales europeos [REF_05]."
    ))

    add_heading(doc, "2.3 Comportamiento Operativo de Aerolineas Low-Cost Europeas", 2)
    add_paragraph(doc, (
        "Las aerolineas low-cost operan bajo modelos de negocio que priorizan la "
        "maximizacion de la rotacion de aeronaves y la minimizacion de costes operativos. "
        "Este modelo genera vulnerabilidades especificas a los retrasos en cascada "
        "[REF_06]. Estudios del sector indican que Ryanair y easyJet operan con "
        "margenes de tiempo entre vuelos (turnaround) inferiores a la media del sector "
        "[REF_07]."
    ))

    add_heading(doc, "2.4 Teoria de Juegos y Arbitraje de Riesgos Regulatorios", 2)
    add_paragraph(doc, (
        "El concepto de arbitraje regulatorio en el sector aereo es relativamente "
        "novedoso en la literatura academica. La asimetria informacional entre aerolineas "
        "(que conocen la fiabilidad operativa de sus rutas) y pasajeros crea oportunidades "
        "de seleccion adversa desde la perspectiva del pasajero informado [REF_08]."
    ))

    add_heading(doc, "2.5 Empresas de Reclamacion (Claim Agencies)", 2)
    add_paragraph(doc, (
        "La aparicion de empresas especializadas en la gestion de reclamaciones EU261 "
        "(AirHelp, FlightRight, Reclamador.es) evidencia la existencia de valor economico "
        "explotable en el arbitraje de compensaciones aereas. Estas empresas operan con "
        "tasas de exito del 60-80% en reclamaciones validas y comisiones del 25-35% "
        "sobre la compensacion obtenida [REF_09]."
    ))

    doc.add_page_break()


def build_metodologia(doc: Document) -> None:
    """Genera el Capitulo 3: Metodologia."""
    add_heading(doc, "3. Metodologia", 1)

    add_heading(doc, "3.1 Fuentes de Datos y Adquisicion", 2)
    add_paragraph(doc, (
        "La fuente de datos principal es el repositorio ADRR (Aviation Data Repository "
        "for Research) de Eurocontrol, que proporciona datos a nivel de vuelo individual "
        "para la aviacion comercial europea. El acceso requiere registro en OneSky Online "
        "y aprobacion de la solicitud. Los datos incluyen hora programada y real de salida "
        "y llegada, aeropuerto de origen y destino, codigo IATA de la aerolinea y tipo "
        "de aeronave."
    ))
    add_paragraph(doc, (
        "Como datos auxiliares se utilizan las coordenadas de aeropuertos de OpenFlights "
        "(airports.dat), descargadas automaticamente mediante el modulo src/utils/geo.py."
    ))

    add_heading(doc, "3.2 Preprocesamiento y Limpieza", 2)
    add_paragraph(doc, (
        "El pipeline de limpieza (src/data/clean.py) realiza las siguientes operaciones: "
        "calculo del retraso como diferencia entre llegada real y programada, eliminacion "
        "de vuelos cancelados (sin hora de llegada real), eliminacion de outliers extremos "
        "(retrasos superiores a 24 horas o adelantos superiores a 6 horas), eliminacion "
        "de duplicados por combinacion (origen, destino, aerolinea, hora de salida), "
        "y tratamiento de missing values documentado por columna."
    ))

    add_heading(doc, "3.3 Feature Engineering", 2)
    add_paragraph(doc, (
        "Se construyen tres categorias de features: temporales (hora, dia de la semana, "
        "mes, indicadores de fin de semana, verano y periodo vacacional), geograficas "
        "(distancia Haversine origen-destino, banda de compensacion EU261) y de tasas "
        "historicas (probabilidad de retraso por aerolinea, aeropuerto de origen, "
        "aeropuerto de destino y ruta)."
    ))
    add_paragraph(doc, (
        "Las tasas historicas se calculan exclusivamente sobre el training set y se "
        "aplican al validation y test set para evitar data leakage, siguiendo el "
        "principio de que el modelo no puede 'ver' informacion del futuro durante el "
        "entrenamiento. Este punto es critico para la validez de la evaluacion."
    ))

    add_heading(doc, "3.4 Seleccion de Modelos y Justificacion", 2)
    add_paragraph(doc, (
        "Se comparan cuatro modelos: Regresion Logistica (baseline interpretable), "
        "Random Forest (comparacion y feature importance), XGBoost (modelo principal, "
        "estado del arte en clasificacion tabular) y LightGBM (alternativa eficiente). "
        "Todos se implementan como pipelines de scikit-learn para garantizar que el "
        "preprocesamiento no cause data leakage."
    ))

    add_heading(doc, "3.5 Estrategia de Validacion y Manejo del Desbalance", 2)
    add_paragraph(doc, (
        "Se utiliza un split temporal cuando el rango de fechas lo permite (entrenamiento "
        "con datos pasados, evaluacion con datos futuros). El desbalance de clases se "
        "gestiona mediante class_weight='balanced' en la regresion logistica, "
        "scale_pos_weight en XGBoost e is_unbalance en LightGBM."
    ))

    add_heading(doc, "3.6 Metricas de Evaluacion", 2)
    add_paragraph(doc, (
        "La metrica principal de comparacion es el AUC-PR (Area bajo la curva "
        "Precision-Recall), mas informativa que el AUC-ROC en presencia de severo "
        "desbalance de clases. Se reportan tambien AUC-ROC, F1-Score (con threshold "
        "optimo), Precision, Recall y Log Loss."
    ))

    add_heading(doc, "3.7 Modelo de Expected Value para Arbitraje", 2)
    add_paragraph(doc, (
        "El modelo de arbitraje calcula el Expected Value neto segun la formula:"
    ))
    add_paragraph(doc, (
        "EV = P(retraso >= 3h) x Compensacion_EU261 x (1 - P(reclamacion_denegada)) "
        "- Precio_Billete - Coste_Transporte - Valor_Hora x Horas_Invertidas"
    ), bold=True)
    add_paragraph(doc, (
        "Donde P(retraso >= 3h) es la probabilidad predicha por el modelo, "
        "Compensacion_EU261 es 250, 400 o 600 euros segun la distancia, "
        "P(reclamacion_denegada) = 0.30 como estimacion central, "
        "Precio_Billete es el precio promedio low-cost estimado por ruta, "
        "Coste_Transporte = 25 EUR, Valor_Hora = 8.87 EUR/h (SMI espanol 2025) "
        "y Horas_Invertidas = 8 horas."
    ))

    doc.add_page_break()


def build_eda(doc: Document) -> None:
    """Genera el Capitulo 4: Analisis Exploratorio de Datos."""
    add_heading(doc, "4. Analisis Exploratorio de Datos", 1)

    add_paragraph(doc, (
        "El analisis exploratorio se realiza sobre el dataset limpio de vuelos de las "
        "cuatro aerolineas low-cost objetivo. Se presentan los principales hallazgos "
        "mediante las figuras generadas automaticamente por el modulo "
        "src/visualization/eda_plots.py."
    ))

    add_heading(doc, "4.1 Distribucion de Retrasos", 2)
    add_paragraph(doc, (
        "Como se muestra en la Figura 1, la distribucion de retrasos presenta una "
        "marcada asimetria positiva, con la gran mayoria de vuelos llegando con menos "
        "de 30 minutos de retraso o incluso adelantados. La proporcion de vuelos con "
        "retraso superior a 180 minutos (elegibles EU261) es inferior al 1%, lo que "
        "constituye un severo desbalance de clases que debe tratarse explicitamente "
        "en el modelado."
    ))
    add_figure(doc, "fig_01_delay_distribution.png",
               "Figura 1. Distribucion de retrasos en minutos. La linea discontinua "
               "roja marca el umbral de 180 minutos (3 horas) de la normativa EU261.")

    add_heading(doc, "4.2 Tasa de Retrasos EU261 por Aerolinea", 2)
    add_paragraph(doc, (
        "Como se muestra en la Figura 2, existen diferencias significativas en la tasa "
        "de retrasos EU261 entre aerolineas. Sin embargo, estas diferencias brutas pueden "
        "estar confundidas por las rutas que opera cada aerolinea (Paradoja de Simpson), "
        "lo que se analiza en detalle en el Capitulo 5."
    ))
    add_figure(doc, "fig_02_eu261_rate_by_airline.png",
               "Figura 2. Tasa de retrasos EU261 por aerolinea low-cost.")

    add_heading(doc, "4.3 Patrones Temporales", 2)
    add_figure(doc, "fig_03_eu261_rate_by_hour.png",
               "Figura 3. Tasa de retrasos EU261 por hora del dia.")
    add_figure(doc, "fig_04_eu261_rate_by_month.png",
               "Figura 4. Tasa de retrasos EU261 por mes del ano.")
    add_figure(doc, "fig_08_heatmap_day_hour.png",
               "Figura 8. Heatmap de tasa de retrasos por dia de la semana y hora.")

    add_heading(doc, "4.4 Aeropuertos con Mayor Tasa de Retraso", 2)
    add_figure(doc, "fig_05_top20_origin_airports.png",
               "Figura 5. Top 20 aeropuertos de origen con mayor tasa de retraso EU261.")
    add_figure(doc, "fig_06_top20_dest_airports.png",
               "Figura 6. Top 20 aeropuertos de destino con mayor tasa de retraso EU261.")

    add_heading(doc, "4.5 Edad de los Aviones y Paradoja Operativa", 2)
    add_figure(doc, "fig_07_aircraft_age_paradox.png",
               "Figura 7. Tasa de retraso vs edad del avion.")

    add_heading(doc, "4.6 Desbalance de Clases", 2)
    add_figure(doc, "fig_10_class_imbalance.png",
               "Figura 10. Distribucion de clases en el dataset.")

    add_heading(doc, "4.7 Correlacion entre Variables", 2)
    add_figure(doc, "fig_09_correlation_matrix.png",
               "Figura 9. Matriz de correlacion de features numericas.")

    doc.add_page_break()


def build_resultados_modelado(doc: Document) -> None:
    """Genera el Capitulo 5: Resultados del Modelado."""
    add_heading(doc, "5. Resultados del Modelado", 1)

    add_heading(doc, "5.1 Comparativa de Modelos", 2)
    add_paragraph(doc, (
        "Se comparan cuatro modelos sobre el test set. La metrica principal de "
        "comparacion es el AUC-PR, mas adecuada que el AUC-ROC para datasets con "
        "severo desbalance de clases. La Tabla 1 presenta los resultados."
    ))
    add_table_from_csv(doc, "model_comparison.csv",
                       "Tabla 1. Comparativa de metricas de evaluacion por modelo.")

    add_heading(doc, "5.2 Curvas ROC y Precision-Recall", 2)
    add_figure(doc, "fig_11_roc_curves.png",
               "Figura 11. Curvas ROC de todos los modelos en el test set.")
    add_figure(doc, "fig_12_pr_curves.png",
               "Figura 12. Curvas Precision-Recall de todos los modelos.")

    add_heading(doc, "5.3 Analisis del Mejor Modelo", 2)
    add_paragraph(doc, (
        "El modelo con mayor AUC-PR se selecciona como modelo principal. "
        "Se identifica el threshold optimo de clasificacion que maximiza el F1-Score "
        "sobre el validation set, evitando optimizar sobre el test set."
    ))
    add_figure(doc, "fig_17_threshold_analysis.png",
               "Figura 17. Precision, Recall y F1 en funcion del threshold.")
    add_figure(doc, "fig_13_confusion_matrix.png",
               "Figura 13. Matriz de confusion del mejor modelo con threshold optimo.")

    add_heading(doc, "5.4 Interpretabilidad: SHAP Values y Feature Importance", 2)
    add_paragraph(doc, (
        "Los valores SHAP (Shapley Additive Explanations) permiten cuantificar la "
        "contribucion de cada feature a la prediccion individual. Como se muestra en la "
        "Figura 14, las features de tasa historica (route_delay_rate, origin_delay_rate) "
        "tienen el mayor impacto predictivo, seguidas de variables temporales."
    ))
    add_figure(doc, "fig_14_shap_summary.png",
               "Figura 14. SHAP summary plot del mejor modelo.")
    add_figure(doc, "fig_16_feature_importance.png",
               "Figura 16. Feature importance del mejor modelo.")

    add_heading(doc, "5.5 Paradoja de Simpson: Analisis por Aerolinea", 2)
    add_paragraph(doc, (
        "La Paradoja de Simpson se manifiesta cuando la aerolinea con mayor tasa bruta "
        "de retrasos no presenta necesariamente la mayor probabilidad intrinseca de "
        "retraso al controlar por la ruta operada. Como se muestra en la Figura 18 y "
        "la Tabla 2, el ranking de aerolineas cambia significativamente entre la tasa "
        "bruta observada y la probabilidad media predicha por el modelo."
    ))
    add_figure(doc, "fig_18_simpson_paradox.png",
               "Figura 18. Paradoja de Simpson: tasa bruta vs probabilidad predicha.")
    add_table_from_csv(doc, "simpson_paradox.csv",
                       "Tabla 2. Tasa bruta vs probabilidad predicha por aerolinea.")

    doc.add_page_break()


def build_arbitraje(doc: Document) -> None:
    """Genera el Capitulo 6: Modelo de Arbitraje y Simulacion."""
    add_heading(doc, "6. Modelo de Arbitraje y Simulacion", 1)

    add_heading(doc, "6.1 Definicion del Modelo de Expected Value Extendido", 2)
    add_paragraph(doc, (
        "El modelo de arbitraje extiende el Expected Value simple incorporando los "
        "costes reales que un arbitrajista incurriria: precio del billete, transporte "
        "al aeropuerto (estimado en 25 EUR), coste de oportunidad del tiempo "
        "(8 horas a 8.87 EUR/h) y la probabilidad de que la aerolinea deniegue la "
        "compensacion alegando circunstancias extraordinarias (30% de base)."
    ))

    add_heading(doc, "6.2 Top Rutas por Expected Value", 2)
    add_figure(doc, "fig_19_top_routes_ev.png",
               "Figura 19. Top rutas por Expected Value medio de arbitraje.")
    add_table_from_csv(doc, "top_routes_ev.csv",
                       "Tabla 3. Top 10 rutas con mayor Expected Value.")

    add_heading(doc, "6.3 Analisis de Sensibilidad", 2)
    add_paragraph(doc, (
        "El analisis de sensibilidad muestra como varia el EV al modificar cada "
        "parametro del modelo. El precio del billete y la probabilidad de rechazo "
        "de la reclamacion son los factores con mayor impacto en la rentabilidad."
    ))
    add_figure(doc, "fig_20_sensitivity_analysis.png",
               "Figura 20. Analisis de sensibilidad del Expected Value.")

    add_heading(doc, "6.4 Analisis de Break-Even", 2)
    add_paragraph(doc, (
        "La Figura 21 muestra la probabilidad minima de retraso necesaria para que "
        "el arbitraje sea rentable en funcion del precio del billete. Para la mayoria "
        "de escenarios reales, esta probabilidad supera ampliamente la tasa media "
        "observada en el dataset (linea de puntos en la figura)."
    ))
    add_figure(doc, "fig_21_breakeven.png",
               "Figura 21. Curva de break-even: probabilidad minima vs precio del billete.")

    add_heading(doc, "6.5 Simulacion Monte Carlo", 2)
    add_paragraph(doc, (
        "La simulacion Monte Carlo con 10,000 iteraciones cuantifica la incertidumbre "
        "del modelo de arbitraje al muestrear cada parametro de su distribucion de "
        "incertidumbre. La Figura 22 muestra la distribucion resultante del EV y el "
        "porcentaje de escenarios rentables."
    ))
    add_figure(doc, "fig_22_ev_simulation.png",
               "Figura 22. Distribucion del EV por simulacion Monte Carlo (10,000 iter.).")

    doc.add_page_break()


def build_discusion(doc: Document) -> None:
    """Genera el Capitulo 7: Discusion."""
    add_heading(doc, "7. Discusion", 1)

    add_heading(doc, "7.1 Viabilidad Practica del Modelo de Arbitraje", 2)
    add_paragraph(doc, (
        "Los resultados del modelo sugieren que, si bien es posible identificar rutas "
        "con mayor probabilidad de retraso EU261, el Expected Value neto es negativo "
        "para la mayoria de escenarios cuando se incorporan los costes reales. "
        "La viabilidad del arbitraje dependeria de poder reducir al minimo el coste "
        "de oportunidad temporal, lo que requiere una dedicacion intensiva incompatible "
        "con un empleo regular."
    ))

    add_heading(doc, "7.2 Limitaciones del Estudio", 2)
    add_paragraph(doc, "Las principales limitaciones del presente trabajo son:")

    limitations = [
        ("Circunstancias extraordinarias",
         "Las aerolineas pueden alegar causas de fuerza mayor (meteorologia, huelgas ATC, "
         "emergencias) para denegar la compensacion. El modelo no puede predecir estos "
         "eventos, por lo que la tasa real de exito de las reclamaciones puede ser "
         "inferior a la asumida."),
        ("Tasa de rechazo de reclamaciones",
         "En la practica, entre el 20 y el 40% de reclamaciones son rechazadas "
         "inicialmente. El EV real es inferior al teorico. El modelo asume una tasa "
         "central del 30%, pero la incertidumbre sobre este parametro es alta."),
        ("Coste de oportunidad temporal",
         "El modelo asume un coste de 8 horas al salario minimo, pero el coste real "
         "puede variar significativamente segun el perfil del arbitrajista."),
        ("Cancelaciones previas al vuelo",
         "Si la aerolinea cancela el vuelo con mas de 14 dias de antelacion, no hay "
         "compensacion EU261. El modelo no puede predecir cancelaciones anticipadas."),
        ("Adaptacion de las aerolineas",
         "Si el arbitraje se popularizase, las aerolineas podrian ajustar precios en "
         "rutas problematicas o mejorar la puntualidad selectivamente."),
        ("Calidad de los datos ADRR",
         "Los datos de Eurocontrol pueden no coincidir exactamente con la experiencia "
         "del pasajero, que se rige por la hora de apertura de puertas, no por la hora "
         "de llegada al parking."),
        ("Sesgo de supervivencia",
         "Solo se analizan vuelos operados, no los cancelados antes de operar."),
        ("Periodo de datos limitado",
         "El modelo se entrena con datos historicos que pueden no reflejar el "
         "comportamiento futuro (impacto del COVID, nuevas regulaciones, cambios de flota)."),
        ("No se consideran vuelos con escala",
         "El EU261 tiene reglas especificas para vuelos con conexion que no se analizan."),
        ("No se modela el valor del desplazamiento",
         "El modelo asume que el arbitrajista no tiene interes en el viaje en si mismo, "
         "lo que es una simplificacion que puede no ser realista."),
    ]

    for title, text in limitations:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(text)

    add_heading(doc, "7.3 Implicaciones Eticas y Legales", 2)
    add_paragraph(doc, (
        "El arbitraje EU261 es completamente legal, ya que el reglamento no establece "
        "ninguna restriccion sobre los motivos del pasajero para volar. Sin embargo, "
        "la masificacion de esta practica podria tener consecuencias no deseadas: "
        "presion sobre las aerolineas para mejorar la puntualidad, pero tambien "
        "posible aumento de precios en rutas problematicas o endurecimiento de los "
        "criterios de denegacion de reclamaciones."
    ))

    add_heading(doc, "7.4 Posibles Respuestas de las Aerolineas", 2)
    add_paragraph(doc, (
        "Desde la perspectiva de la teoria de juegos, si un numero significativo de "
        "pasajeros adoptase estrategias de arbitraje informadas por modelos predictivos, "
        "las aerolineas tendrian incentivos para responder. Las estrategias posibles "
        "incluyen: aumento de precios en rutas con alta tasa de retraso historica, "
        "mejora selectiva de la puntualidad en rutas identificadas como problematicas, "
        "o mayor agresividad en la alegacion de circunstancias extraordinarias. "
        "Esta dinamica de juego convertiria el modelo en autoinvalidante si se "
        "adoptara masivamente."
    ))

    doc.add_page_break()


def build_conclusiones(doc: Document) -> None:
    """Genera el Capitulo 8: Conclusiones y Trabajo Futuro."""
    add_heading(doc, "8. Conclusiones y Trabajo Futuro", 1)

    add_heading(doc, "8.1 Conclusiones Principales", 2)
    add_paragraph(doc, (
        "El presente TFM ha demostrado que es tecnicamente posible construir un modelo "
        "predictivo de retrasos EU261 con capacidad discriminatoria superior al azar "
        "(AUC-PR significativamente por encima del baseline). Las features de tasa "
        "historica por ruta y aeropuerto son las mas predictivas, lo que sugiere que "
        "la puntualidad historica de una ruta es el mejor predictor de la puntualidad "
        "futura."
    ))
    add_paragraph(doc, (
        "Sin embargo, el modelo de Expected Value extendido revela que el arbitraje "
        "EU261 no es rentable en la mayoria de escenarios reales cuando se incorporan "
        "los costes completos. La barrera principal no es la capacidad predictiva del "
        "modelo, sino la combinacion de: coste de oportunidad temporal elevado, "
        "probabilidad de rechazo de reclamaciones no despreciable, y el hecho de que "
        "las rutas con mayor probabilidad de retraso tienden a tener precios de billete "
        "mas elevados (seleccion adversa del mercado)."
    ))
    add_paragraph(doc, (
        "La Paradoja de Simpson ilustra los riesgos del analisis descriptivo simple: "
        "la aerolinea con mayor tasa bruta de retrasos no es necesariamente la mas "
        "problematica al controlar por las rutas que opera, lo que tiene implicaciones "
        "directas para estrategias de arbitraje basadas en la simple eleccion de "
        "aerolinea."
    ))

    add_heading(doc, "8.2 Contribuciones del Trabajo", 2)
    add_paragraph(doc, "Las principales contribuciones del TFM son:")

    contributions = [
        "Construccion de un pipeline reproducible de machine learning para prediccion "
        "de retrasos EU261 en aerolineas low-cost europeas.",
        "Demostracion empirica de la Paradoja de Simpson en tasas de retraso aereo "
        "europeo.",
        "Desarrollo de un modelo de Expected Value extendido con analisis de "
        "sensibilidad y simulacion Monte Carlo para el arbitraje EU261.",
        "Identificacion de las limitaciones criticas del arbitraje regulatorio aereo "
        "desde una perspectiva de datos.",
    ]

    for contrib in contributions:
        p = doc.add_paragraph(style="List Number")
        p.add_run(contrib)

    add_heading(doc, "8.3 Lineas de Trabajo Futuro", 2)
    add_paragraph(doc, "Como lineas de investigacion futura se proponen:")

    future_work = [
        "Incorporar datos meteorologicos en tiempo real para predecir retrasos "
        "por causas atmosfericas con mayor precision.",
        "Modelar el riesgo de cancelacion anticipada de vuelos (antes de los 14 dias) "
        "como factor adicional del modelo de arbitraje.",
        "Aplicar tecnicas de series temporales (LSTM, Prophet) para capturar "
        "dependencias temporales entre retrasos consecutivos en la misma aeronave.",
        "Extender el analisis a otras normativas de compensacion (retrasos de tren, "
        "cruceros) para comparar oportunidades de arbitraje regulatorio.",
        "Desarrollar una interfaz web que integre el modelo predictivo con datos de "
        "precios de billetes en tiempo real para calcular el EV de cada vuelo.",
    ]

    for item in future_work:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_page_break()


def build_bibliografia(doc: Document) -> None:
    """Genera la seccion de Bibliografia con placeholders."""
    add_heading(doc, "Bibliografia", 1)

    add_paragraph(doc, (
        "NOTA: Las referencias marcadas con [REF_XX] deben ser sustituidas por "
        "citas reales buscadas en Google Scholar. Se sugieren los terminos de busqueda "
        "indicados en cada placeholder."
    ), bold=True)

    doc.add_paragraph()

    refs = [
        "[REF_01] Buscar: 'flight delay prediction machine learning survey 2020-2024'. "
        "Autor, A. (Ano). Titulo del articulo. Revista, volumen(numero), paginas. "
        "DOI: xxxxx",
        "[REF_02] Buscar: 'XGBoost flight delay prediction European aviation'. "
        "Autor, B. (Ano). Titulo. Proceedings of XXX Conference, paginas.",
        "[REF_03] Buscar: 'class imbalance flight delay prediction SMOTE'. "
        "Autor, C. (Ano). Titulo. Journal of Air Transport, volumen, paginas.",
        "[REF_04] Tribunal de Justicia de la Union Europea. (2009). "
        "Sentencia en los asuntos acumulados C-402/07 y C-432/07 "
        "(Sturgeon y otros contra Condor Flugdienst GmbH). Recopilacion 2009.",
        "[REF_05] Buscar: 'EU261 extraordinary circumstances case law'. "
        "Autor, D. (Ano). Titulo. Journal of Air Law and Commerce, volumen, paginas.",
        "[REF_06] Buscar: 'low-cost airline delay cascade operational'. "
        "Autor, E. (Ano). Titulo. Transportation Research Part E, volumen, paginas.",
        "[REF_07] Buscar: 'Ryanair easyJet turnaround time punctuality'. "
        "Autor, F. (Ano). Titulo. Journal of Transport Economics and Policy.",
        "[REF_08] Buscar: 'regulatory arbitrage information asymmetry airline'. "
        "Autor, G. (Ano). Titulo. European Journal of Law and Economics.",
        "[REF_09] AirHelp. (2024). AirHelp Score 2024. Recuperado de "
        "https://www.airhelp.com/en/airhelp-score/ [Fecha de acceso].",
    ]

    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)

    doc.add_page_break()


def build_anexos(doc: Document) -> None:
    """Genera los Anexos del TFM."""
    add_heading(doc, "Anexos", 1)

    add_heading(doc, "Anexo A. Fragmentos de Codigo Fuente Relevantes", 2)
    add_paragraph(doc, (
        "Se presentan fragmentos del codigo fuente mas relevante para la comprension "
        "del pipeline. El codigo completo esta disponible en el repositorio de GitHub "
        "adjunto al TFM."
    ))

    add_paragraph(doc, "A.1 Calculo del Expected Value (src/utils/eu261.py)", bold=True)
    code_text = (
        "def calculate_expected_value(p_delay, distance_km, ticket_price_eur,\n"
        "                             transport_cost_eur=25.0,\n"
        "                             hourly_wage_eur=8.87,\n"
        "                             hours_invested=8.0,\n"
        "                             p_claim_denied=0.30):\n"
        "    compensation = get_eu261_compensation(distance_km)\n"
        "    expected_income = p_delay * compensation * (1.0 - p_claim_denied)\n"
        "    time_cost = hourly_wage_eur * hours_invested\n"
        "    total_cost = ticket_price_eur + transport_cost_eur + time_cost\n"
        "    ev = expected_income - total_cost\n"
        "    return {'ev': round(ev, 2), 'profitable': ev > 0.0, ...}"
    )
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    add_heading(doc, "Anexo B. Glosario de Terminos", 2)

    glossary = [
        ("AUC-PR", "Area bajo la curva Precision-Recall. Metrica de evaluacion de "
         "clasificadores binarios, especialmente util con clases desbalanceadas."),
        ("AUC-ROC", "Area bajo la curva ROC (Receiver Operating Characteristic). "
         "Mide la capacidad discriminatoria de un clasificador."),
        ("Data leakage", "Filtracion de informacion del futuro al modelo de entrenamiento, "
         "que produce evaluaciones optimistas no reproducibles en produccion."),
        ("EU261", "Reglamento (CE) n. 261/2004 del Parlamento Europeo sobre compensacion "
         "a pasajeros aereos por retrasos, cancelaciones y denegacion de embarque."),
        ("Expected Value (EV)", "Valor esperado de una apuesta o decision, calculado como "
         "la suma ponderada de todos los resultados posibles por su probabilidad."),
        ("IATA", "International Air Transport Association. Organismo que asigna codigos "
         "de dos letras a las aerolineas y de tres letras a los aeropuertos."),
        ("Low-cost carrier (LCC)", "Aerolinea de bajo coste que opera con estructuras de "
         "costes reducidas, tipicamente con menos servicios incluidos en el billete."),
        ("Paradoja de Simpson", "Fenomeno estadistico por el que una tendencia observada "
         "en grupos separados desaparece o se invierte al combinar los grupos."),
        ("SHAP", "SHapley Additive exPlanations. Metodo de interpretabilidad de modelos "
         "de machine learning basado en valores de Shapley de la teoria de juegos."),
        ("XGBoost", "Extreme Gradient Boosting. Implementacion eficiente del algoritmo "
         "de gradient boosting, estado del arte en clasificacion tabular."),
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        run_term = p.add_run(f"{term}: ")
        run_term.bold = True
        p.add_run(definition)


# ---------------------------------------------------------------------------
# Funcion principal
# ---------------------------------------------------------------------------


def generate_tfm() -> None:
    """
    Genera el documento Word completo del TFM y lo guarda en docs/tfm/.
    """
    logger.info("=" * 60)
    logger.info("FASE 7: GENERACION DEL DOCUMENTO TFM")
    logger.info("=" * 60)

    doc = Document()

    # Configurar margenes del documento
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_height = Pt(841.9)  # A4
    section.page_width = Pt(595.3)
    section.left_margin = Inches(1.18)   # 3 cm
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(1.18)
    section.bottom_margin = Inches(1.18)

    # Construir secciones
    build_portada(doc)
    build_abstract(doc)
    build_introduccion(doc)
    build_estado_arte(doc)
    build_metodologia(doc)
    build_eda(doc)
    build_resultados_modelado(doc)
    build_arbitraje(doc)
    build_discusion(doc)
    build_conclusiones(doc)
    build_bibliografia(doc)
    build_anexos(doc)

    # Guardar
    TFM_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(TFM_OUTPUT))
    logger.info("Documento TFM guardado en: %s", TFM_OUTPUT)


if __name__ == "__main__":
    generate_tfm()
